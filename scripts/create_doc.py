"""
CyberForge - Script de création du document Word de documentation
Author: KAOUANE WALID
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = r"C:\Users\Walid\Desktop\projet github 1\CyberForge_Documentation.docx"

# ── Couleurs ──────────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1F, 0x39, 0x64)   # titres principaux
MID_BLUE   = RGBColor(0x2E, 0x75, 0xB6)   # sous-titres / accents
LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)   # fond en-tête tableau
RED        = RGBColor(0xC0, 0x00, 0x00)   # alerte / warning
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x00, 0x00, 0x00)
CODE_BG    = RGBColor(0x1E, 0x1E, 0x1E)   # fond blocs de code
CODE_FG    = RGBColor(0xD4, 0xD4, 0xD4)   # texte blocs de code


def set_cell_bg(cell, hex_color: str):
    """Applique une couleur de fond à une cellule de tableau."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="CCCCCC"):
    """Ajoute des bordures légères à une cellule."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_footer(doc):
    """Ajoute un pied de page sur toutes les pages."""
    section = doc.sections[0]
    footer  = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CyberForge — Documentation Technique  |  KAOUANE WALID  |  2026")
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    run.font.name  = "Arial"

    # Ligne de séparation au-dessus du footer
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top  = OxmlElement("w:top")
    top.set(qn("w:val"),   "single")
    top.set(qn("w:sz"),    "6")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "2E75B6")
    pBdr.append(top)
    pPr.append(pBdr)


def add_heading(doc, text, level=1):
    """Ajoute un titre avec mise en forme personnalisée."""
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    if level == 1:
        run.font.size  = Pt(18)
        run.font.color.rgb = DARK_BLUE
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
        # Ligne de soulignement
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "2E75B6")
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size  = Pt(14)
        run.font.color.rgb = MID_BLUE
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
    elif level == 3:
        run.font.size  = Pt(12)
        run.font.color.rgb = DARK_BLUE
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
    return p


def add_body(doc, text, bold_parts=None):
    """Ajoute un paragraphe de corps de texte."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    # Gestion du gras inline **texte**
    import re
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        if i % 2 == 1:  # entre les **
            run.bold = True
    return p


def add_numbered_item(doc, number, text):
    """Ajoute un élément numéroté."""
    p   = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return p


def add_bullet(doc, text):
    """Ajoute un élément à puce."""
    p = doc.add_paragraph(style="List Bullet")
    import re
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        if i % 2 == 1:
            run.bold = True
    return p


def add_code_block(doc, code_text):
    """Ajoute un bloc de code stylisé (fond sombre)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    # Fond gris clair via ombrage
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "1E1E1E")
    pPr.append(shd)
    run = p.add_run(code_text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(9)
    run.font.color.rgb = CODE_FG
    return p


def add_table(doc, headers, rows, col_widths=None):
    """Ajoute un tableau avec en-têtes colorés."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # En-tête
    hdr_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = header
        set_cell_bg(cell, "1F3964")
        set_cell_borders(cell, "3A5990")
        run = cell.paragraphs[0].runs[0]
        run.bold       = True
        run.font.color.rgb = WHITE
        run.font.name  = "Arial"
        run.font.size  = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Lignes de données
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg  = "EAF0FB" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = cell_text
            set_cell_bg(cell, bg)
            set_cell_borders(cell, "CCCCCC")
            run = cell.paragraphs[0].runs[0]
            run.font.name = "Arial"
            run.font.size = Pt(10)

    # Largeurs de colonnes
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[i])

    doc.add_paragraph()  # espacement après tableau


# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT PRINCIPAL
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# Marges
section = doc.sections[0]
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

add_footer(doc)

# ─── PAGE DE TITRE ─────────────────────────────────────────────────────────

# Espace vertical
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CyberForge")
run.bold = True
run.font.name  = "Arial"
run.font.size  = Pt(40)
run.font.color.rgb = DARK_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Documentation Technique du Projet")
run.font.name  = "Arial"
run.font.size  = Pt(20)
run.font.color.rgb = MID_BLUE

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Laboratoire de Cybersécurité Personnel")
run.font.name  = "Arial"
run.font.size  = Pt(14)
run.font.color.rgb = RGBColor(0x50, 0x50, 0x50)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Auteur : KAOUANE WALID")
run.bold = True
run.font.name  = "Arial"
run.font.size  = Pt(13)
run.font.color.rgb = DARK_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("2026")
run.font.name  = "Arial"
run.font.size  = Pt(12)
run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

# Saut de page
doc.add_page_break()

# ─── 1. PRÉSENTATION DU PROJET ──────────────────────────────────────────────

add_heading(doc, "1.  Présentation du Projet", 1)
add_body(doc, (
    "CyberForge est un laboratoire virtuel personnel conçu pour simuler des scénarios "
    "d'intrusion et de réponse réalistes dans un environnement Windows/Active Directory "
    "isolé. Le projet démontre des capacités de cybersécurité offensives (Red Team) et "
    "défensives (Blue Team)."
))
add_body(doc, (
    "Ce laboratoire a été réalisé en **10 jours** dans le but d'acquérir une expérience "
    "pratique avec des outils SOC (Security Operations Center) standards de l'industrie, "
    "couvrant la détection des menaces, la réponse automatisée aux incidents et la "
    "remédiation des logiciels malveillants."
))
add_body(doc, "**Auteur :** KAOUANE WALID")
add_body(doc, "**Durée de réalisation :** 10 jours")
add_body(doc, "**Domaine :** Cybersécurité — Red Team / Blue Team")

# ─── 2. ARCHITECTURE ────────────────────────────────────────────────────────

add_heading(doc, "2.  Architecture du Laboratoire", 1)
add_heading(doc, "2.1  Machines Virtuelles", 2)

add_table(doc,
    headers=["Machine Virtuelle", "Système d'Exploitation", "Rôle", "RAM"],
    rows=[
        ["Wazuh Server",   "Ubuntu 22.04 LTS",       "Gestionnaire SIEM / EDR",       "4 Go"],
        ["Windows Server", "Windows Server 2019",    "Contrôleur AD + Cible",          "4 Go"],
        ["Kali Linux",     "Kali Linux 2024",        "Machine attaquante (Red Team)",  "2 Go"],
    ],
    col_widths=[1.5, 2.0, 2.5, 0.8]
)

add_heading(doc, "2.2  Outils de Sécurité Utilisés", 2)

add_table(doc,
    headers=["Outil", "Rôle", "Description"],
    rows=[
        ["Wazuh",          "SIEM / EDR",         "Plateforme open-source de détection des menaces, surveillance d'intégrité et réponse aux incidents"],
        ["Suricata",       "IDS Réseau",          "Analyse du trafic réseau en temps réel et détection d'intrusions"],
        ["VirusTotal",     "Analyse de Malwares", "Analyse de malwares basée sur le cloud, intégrée à Wazuh FIM"],
        ["Kali Linux / Nmap", "Reconnaissance",  "Outils de sécurité offensifs pour la simulation Red Team"],
        ["Hydra",          "Brute Force",         "Outil de test de force brute pour simulation d'attaques RDP"],
    ],
    col_widths=[1.6, 1.5, 3.7]
)

# ─── 3. SCÉNARIO 1 ──────────────────────────────────────────────────────────

add_heading(doc, "3.  Scénario 1 — Détection de la Reconnaissance Réseau", 1)

add_heading(doc, "3.1  Objectif", 2)
add_body(doc, "Détecter et alerter sur les activités de balayage réseau initiées depuis la machine Kali Linux.")

add_heading(doc, "3.2  Outils Impliqués", 2)
add_bullet(doc, "**Attaque :** Nmap (depuis Kali Linux)")
add_bullet(doc, "**Détection :** Suricata IDS")
add_bullet(doc, "**Corrélation & Alerte :** Wazuh SIEM")

add_heading(doc, "3.3  Déroulement", 2)
add_numbered_item(doc, 1, "L'attaquant lance un scan SYN Nmap depuis Kali Linux vers le réseau cible")
add_numbered_item(doc, 2, "Suricata analyse le trafic réseau en temps réel et identifie les signatures de scan (règles ET SCAN)")
add_numbered_item(doc, 3, "Suricata génère une alerte et l'écrit dans le fichier eve.json")
add_numbered_item(doc, 4, "Wazuh ingère le fichier eve.json de Suricata et corrèle les alertes")
add_numbered_item(doc, 5, "Le tableau de bord Wazuh affiche l'alerte avec les métadonnées complètes des paquets")

add_heading(doc, "3.4  Commande d'Attaque (Kali Linux)", 2)
add_code_block(doc, "# Scan SYN furtif\nnmap -sS -T4 -p 1-1000 <TARGET_IP>\n\n# Détection de services\nnmap -sV -T4 -p 22,80,443,3389 <TARGET_IP>\n\n# Empreinte OS\nsudo nmap -O <TARGET_IP>")

add_heading(doc, "3.5  Règles de Détection Wazuh", 2)
add_bullet(doc, "Règle **100001** (niveau 10) : Suricata détecte une fuite d'informations réseau (ET SCAN)")
add_bullet(doc, "Règle **100002** (niveau 12) : Signature Nmap détectée par Suricata")
add_bullet(doc, "Règle **100003** (niveau 14) : Reconnaissance répétée depuis la même IP (haute fréquence)")

add_heading(doc, "3.6  Résultats", 2)
add_body(doc, (
    "Suricata a identifié avec succès les scans SYN/ACK et les phases de cartographie réseau. "
    "Les alertes ont été remontées au tableau de bord Wazuh avec les métadonnées complètes "
    "des paquets en **moins de 5 secondes**."
))
add_body(doc, "**MITRE ATT&CK :** T1046 — Network Service Discovery")

# ─── 4. SCÉNARIO 2 ──────────────────────────────────────────────────────────

add_heading(doc, "4.  Scénario 2 — Défense contre le Brute Force RDP", 1)

add_heading(doc, "4.1  Objectif", 2)
add_body(doc, "Bloquer automatiquement les adresses IP effectuant des attaques par force brute sur le service RDP.")

add_heading(doc, "4.2  Outils Impliqués", 2)
add_bullet(doc, "**Attaque :** Hydra (depuis Kali Linux)")
add_bullet(doc, "**Détection :** Wazuh — Règle 60204 (événements Windows 4625)")
add_bullet(doc, "**Réponse Active :** Script block-ip.cmd via Windows netsh")

add_heading(doc, "4.3  Déroulement", 2)
add_numbered_item(doc, 1, "L'attaquant lance Hydra depuis Kali Linux pour effectuer un brute force RDP")
add_numbered_item(doc, 2, "Windows enregistre les échecs d'authentification (Event ID 4625, LogonType 10 = RDP)")
add_numbered_item(doc, 3, "L'agent Wazuh transmet les événements au gestionnaire Wazuh")
add_numbered_item(doc, 4, "La règle 60204 se déclenche après plusieurs échecs consécutifs")
add_numbered_item(doc, 5, "Wazuh exécute le script de réponse active block-ip.cmd")
add_numbered_item(doc, 6, "La commande netsh ajoute une règle de pare-feu bloquant l'IP source pendant 600 secondes")
add_numbered_item(doc, 7, "Après 600 secondes, la règle est automatiquement supprimée")

add_heading(doc, "4.4  Commande d'Attaque (Kali Linux)", 2)
add_code_block(doc, "hydra -l Administrator -P /usr/share/wordlists/rockyou.txt \\\n      -t 4 rdp://<TARGET_IP> -V")

add_heading(doc, "4.5  Commande de Blocage (Windows — Active Response)", 2)
add_code_block(doc, 'netsh advfirewall firewall add rule name="WAZUH_BLOCK_[IP]" \\\n    protocol=any dir=in action=block remoteip=[IP_ATTAQUANT]')

add_heading(doc, "4.6  Résultats", 2)
add_body(doc, (
    "La règle Wazuh 60204 déclenche un blocage automatique de l'IP via les commandes Windows netsh "
    "pour **600 secondes (10 minutes)**, neutralisant efficacement l'attaque par force brute en "
    "temps réel. Temps de réponse : **moins de 3 secondes**."
))
add_body(doc, "**MITRE ATT&CK :** T1110.001 — Password Guessing")

# ─── 5. SCÉNARIO 3 ──────────────────────────────────────────────────────────

add_heading(doc, "5.  Scénario 3 — Détection et Suppression de Malware", 1)

add_heading(doc, "5.1  Objectif", 2)
add_body(doc, "Détecter, analyser et supprimer automatiquement les fichiers malveillants déposés sur le système.")

add_heading(doc, "5.2  Outils Impliqués", 2)
add_bullet(doc, "**Fichier de test :** Fichier de test antivirus standard EICAR")
add_bullet(doc, "**Détection :** Wazuh FIM (File Integrity Monitoring)")
add_bullet(doc, "**Analyse :** API VirusTotal")
add_bullet(doc, "**Remédiation :** Exécutable Windows personnalisé delete_malware.exe")

add_heading(doc, "5.3  Déroulement", 2)
add_numbered_item(doc, 1, r"Un fichier EICAR est déposé dans un répertoire surveillé (C:\Users\Public\Downloads)")
add_numbered_item(doc, 2, "Wazuh FIM détecte l'ajout du nouveau fichier en temps réel")
add_numbered_item(doc, 3, "Le hash SHA256 du fichier est soumis automatiquement à l'API VirusTotal")
add_numbered_item(doc, 4, "VirusTotal retourne un verdict malveillant (60+/72 moteurs antivirus)")
add_numbered_item(doc, 5, "Wazuh déclenche la réponse active avec delete_malware.exe")
add_numbered_item(doc, 6, "L'exécutable supprime le fichier et journalise l'action")
add_numbered_item(doc, 7, "Wazuh FIM confirme la suppression du fichier")

add_heading(doc, "5.4  Commande de Dépôt du Fichier Test (Windows PowerShell)", 2)
add_code_block(doc,
    r'$eicar = "X5O!P%@AP[4\PZX54(P^)7CC)7}$EICAR-STANDARD-ANTIVIRUS-TEST-FILE!$H+H*"' + "\n" +
    r'Set-Content -Path "C:\Users\Public\Downloads\test_malware.txt" -Value $eicar'
)

add_heading(doc, "5.5  Résultats VirusTotal", 2)
add_body(doc, (
    "Le fichier de test EICAR a été détecté comme malveillant par **plus de 60 moteurs "
    "antivirus sur 72**, validant l'intégration complète de la chaîne "
    "détection → analyse → remédiation. Temps de réponse global : **moins de 10 secondes**."
))
add_body(doc, "**MITRE ATT&CK :** T1204 — User Execution: Malicious File")

# ─── 6. RÉCAPITULATIF ───────────────────────────────────────────────────────

add_heading(doc, "6.  Récapitulatif des Résultats", 1)

add_table(doc,
    headers=["Scénario", "Attaque Détectée", "Réponse Automatisée", "Temps de Réponse", "Règle Wazuh"],
    rows=[
        ["Reconnaissance Réseau",    "Oui ✓", "Alerte générée",       "< 5 secondes",  "100001 / 100002"],
        ["Brute Force RDP",          "Oui ✓", "IP bloquée (600 s)",   "< 3 secondes",  "60204"],
        ["Dépôt de Malware EICAR",   "Oui ✓", "Fichier supprimé",     "< 10 secondes", "87105"],
    ],
    col_widths=[2.0, 1.4, 1.8, 1.4, 1.2]
)

# ─── 7. ENSEIGNEMENTS ───────────────────────────────────────────────────────

add_heading(doc, "7.  Enseignements Clés", 1)

add_numbered_item(doc, 1, "La corrélation SIEM est critique : les logs individuels signifient peu ; ce sont les patterns qui révèlent les attaques.")
add_numbered_item(doc, 2, "La réponse active comble le fossé entre détection et remédiation sans intervention humaine.")
add_numbered_item(doc, 3, "FIM + renseignement sur les menaces (VirusTotal) crée un pipeline d'analyse automatisé puissant.")
add_numbered_item(doc, 4, "La détection au niveau réseau (Suricata) complète la surveillance basée sur l'hôte (agents Wazuh).")
add_numbered_item(doc, 5, "Les tests avec des fichiers standard (EICAR) permettent de valider la chaîne complète sans risque.")

# ─── 8. AUTEUR ──────────────────────────────────────────────────────────────

add_heading(doc, "8.  Auteur", 1)

p = doc.add_paragraph()
run = p.add_run("KAOUANE WALID")
run.bold = True
run.font.name  = "Arial"
run.font.size  = Pt(14)
run.font.color.rgb = DARK_BLUE

add_body(doc, "Passionné de cybersécurité | Praticien Red Team / Blue Team")
add_body(doc, "Projet réalisé en 10 jours comme laboratoire de cybersécurité pratique personnel.")

# ─── SAUVEGARDE ─────────────────────────────────────────────────────────────

doc.save(OUTPUT_PATH)
print(f"[OK] Document créé : {OUTPUT_PATH}")
